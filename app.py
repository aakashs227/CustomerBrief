# ai_agent.py

import os
import re
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langgraph.prebuilt import create_react_agent
from langchain_core.messages.ai import AIMessage

# Load environment variables
load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Custom system prompt
DEFAULT_PROMPT = (
    "You are an expert business intelligence analyst. When given a company name, provide a detailed report with the following structure:\n\n"
    "1. 🔍 Company Overview\n"
    "2. 💰 Financial Summary (Revenue, Profit, Funding, etc.)\n"
    "3. 📦 Import Activity\n"
    "4. 🚢 Export Activity\n"
    "5. 🌍 Global Presence & Offices\n"
    "6. 🚛 Freight Forwarding History\n"
    "7. 📌 Actionable Insights\n"
    "8. 🔗 List of References (with clickable links)\n\n"
    "i want in detail business overview ,Be clear, and use bullet points or headings for readability. Ensure references are provided within each category itself."
)

def clean_ai_response(response_text: str) -> str:
    """
    Removes the initial disclaimer paragraph from the AI response if it contains fallback indicators.
    """
    paragraphs = response_text.strip().split("\n\n")
    fallback_indicators = [
        "issue retrieving real-time data",
        "based on my expert knowledge",
        "latest publicly available information",
        "compiled from public sources",
        "prior to june 2024",
        "i can provide a report",
        "real-time data isn't available"
    ]
    first_paragraph = paragraphs[0].lower()
    if any(keyword in first_paragraph for keyword in fallback_indicators):
        paragraphs = paragraphs[1:]
    return "\n\n".join(paragraphs).strip()

def get_response_from_ai_agent(llm_id, query, allow_search=False):
    # Define a list of common company suffixes
    company_suffixes = ['Inc', 'Ltd', 'LLC', 'PLC', 'GmbH', 'Industries', 'AG', 'Corp', 'Corporation', 'Co', 'Pvt', 'Limited', 'Group', 'S.A.', 'S.A.S.', 'S.L.', 'S.L.U.']

    suffix_pattern = r'\b(?:[A-Z][a-zA-Z&.\'-]*\s?)+(?:' + '|'.join(company_suffixes) + r')\b'
    capitalized_pattern = r'\b[A-Z][a-zA-Z&.\'-]{2,}\b'

    from company_utils import extract_companies
    potential_companies = extract_companies(query)

    if len(potential_companies) > 1:
        return (
            "⚠️ Important Notice: To ensure clarity and depth in analysis, "
            "our AI system is designed to evaluate one company at a time. "
            "Please revise your query to reference a single organization for a precise and comprehensive report. 🏢"
        )

    llm = ChatOpenAI(model=llm_id)

    # No tools used, as TavilySearch is removed
    agent = create_react_agent(
        model=llm,
        tools=[],  # No external tools
        state_modifier= DEFAULT_PROMPT
    )

    state = {"messages": query}
    response = agent.invoke(state)

    messages = response.get("messages")
    ai_messages = [msg.content for msg in messages if isinstance(msg, AIMessage)]
    return ai_messages[-1] if ai_messages else "No response generated."

# File_operations.py


from io import BytesIO
from docx import Document
import uuid

def generate_docx_file(content: str) -> BytesIO:
    """
    Converts a string (company insights) into a .docx file and returns a BytesIO object.
    """
    doc = Document()
    doc.add_heading('Company Insights', 0)
    
    for line in content.split("\n"):
        doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_share_link(data: str) -> str:
    """
    Mocks generating a shareable link. Replace this with actual logic if needed.
    """
    unique_id = uuid.uuid4().hex
    # In real case, you'd store data in a database or pastebin-like service
    return f"https://mocksharelink.com/report/{unique_id}"


# Company_utils.py



import re

try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
    USE_SPACY = True
except ImportError:
    USE_SPACY = False

company_suffixes = [
    'Inc', 'Ltd', 'LLC', 'PLC', 'GmbH', 'Industries', 'AG', 'Corp',
    'Corporation', 'Co', 'Pvt', 'Limited', 'Group', 'S.A.', 'S.A.S.', 'S.L.', 'S.L.U.'
]

# Suffix-based pattern (case insensitive)
suffix_pattern = r'\b([A-Za-z][\w&.\'-]*(?:\s+[A-Za-z][\w&.\'-]*)*\s+(?:' + '|'.join(company_suffixes) + r'))\b'

# Capitalized phrase pattern (case insensitive)
capitalized_pattern = r'\b(?:[A-Za-z][a-zA-Z&.\'-]+(?:\s+[A-Za-z][a-zA-Z&.\'-]+){0,3})\b'

def extract_companies(text: str):
    companies = set()

    # 1. From suffix-based regex (case insensitive)
    suffix_matches = re.findall(suffix_pattern, text)
    companies.update([m.strip() for m in suffix_matches if m.strip()])

    # 2. From capitalized word phrases (case insensitive)
    capitalized_matches = re.findall(capitalized_pattern, text)
    companies.update([m.strip() for m in capitalized_matches if m.strip()])

    # 3. From spaCy NER if available
    if USE_SPACY:
        doc = nlp(text)
        ner_matches = [ent.text.strip() for ent in doc.ents if ent.label_ == "ORG"]
        companies.update(ner_matches)

    return list(companies)  # Will keep case-sensitive company names intact



# frontend.py
import streamlit as st
import openai
import base64
import requests
from PIL import Image
from io import BytesIO
from docx import Document
import re

# === Setup OpenAI client using new SDK style ===
from openai import OpenAI
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# --- Helper Functions ---
def clean_response(text):
    cleaned_text = re.sub(
        r"(🧠 Company Analysis\s*)(.*?)(?=\n\s*1\.\s*[🔍A-Z])",
        r"\1\n",
        text,
        flags=re.DOTALL
    )
    return cleaned_text.strip()

def slugify(text):
    return "".join(c if c.isalnum() else "_" for c in text.lower()).strip("_")

def generate_docx(query, response):
    doc = Document()
    doc.add_heading("CustomerBrief", level=1)
    doc.add_paragraph(f"Query: {query}")
    doc.add_paragraph("Response:")
    doc.add_paragraph(response)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def contains_multiple_companies(query):
    delimiters = [",", "&", " and ", "/", " vs ", " versus ", ";"]
    count = sum([query.lower().count(d) for d in delimiters])
    return count >= 1


        
 # Show response only once
def show_download_buttons(query, response, key_prefix="main"):
    st.markdown("### 🧠 Company Analysis")

    if contains_multiple_companies(query):
        # Show the warning only once
        if not hasattr(st.session_state, 'multiple_companies_warned'):
            st.warning(
                "⚠️ Important Notice: To ensure clarity and depth in analysis, our AI system is designed to evaluate one company at a time. "
                "Please revise your query to reference a single organization for a precise and comprehensive report. 🏢"
            )
            st.session_state.multiple_companies_warned = True
        

    else:
        file_name = f"{slugify(query)}.docx"
        docx_file = generate_docx(query, response)

        # Top download button
        st.download_button(
            label="📥 Download Analysis",
            data=docx_file,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"{key_prefix}_download_top"
        )

        # Show response
        st.write(response)

        # Bottom download button
        st.download_button(
            label="📥 Download Analysis",
            data=docx_file,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"{key_prefix}_download_bottom"
        )

# --- Streamlit Config ---
st.set_page_config(
    page_title="CustomerBrief |",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Custom CSS ---
st.markdown("""
    <style>
        #MainMenu, header, footer {visibility: hidden;}
        body { background-color: #f4f4f4 !important; }
        section[data-testid="stSidebar"] {background-color: #002b5c;}
        .stButton button {background-color: white !important; color: black !important;}
        .chat-history-item {color: yellow !important; font-size: 16px !important;}
    </style>
""", unsafe_allow_html=True)

# --- Logo Setup ---
logo_url = "https://raw.githubusercontent.com/aakashs227/CustomerBrief/main/WORLDWIDE_Logo%207.png"
try:
    logo_img = Image.open(BytesIO(requests.get(logo_url).content))
    buffered = BytesIO()
    logo_img.save(buffered, format="PNG")
    logo_base64 = base64.b64encode(buffered.getvalue()).decode()
except Exception as e:
    st.error(f"Logo loading error: {e}")
    logo_base64 = ""

# --- Initialize State ---
for key in ["chat_history", "last_query", "selected_menu"]:
    if key not in st.session_state:
        st.session_state[key] = [] if "history" in key else ""

# --- Sidebar ---
with st.sidebar:
    col_logo, _ = st.columns([0.7, 0.3])
    with col_logo:
        st.markdown(
            f"""
            <div style="background-color: white; padding: 3px; display: inline-block; border-radius: 8px;">
                <img src="data:image/png;base64,{logo_base64}" width="150"/>
            </div>
            """, unsafe_allow_html=True
        )
    st.markdown("---")
    st.markdown("<h3 style='color: white;'>Chat History</h3>", unsafe_allow_html=True)

    for i, (q, a) in enumerate(reversed(st.session_state.chat_history)):
        idx = len(st.session_state.chat_history) - 1 - i
        col1, col2 = st.columns([0.85, 0.15])
        with col1:
            st.markdown(f"<div class='chat-history-item'>• {q[:50]}</div>", unsafe_allow_html=True)
        with col2:
            if st.button("⋮", key=f"menu_{i}"):
                st.session_state.selected_menu = idx
        if st.session_state.get("selected_menu") == idx:
            bcols = st.columns(3)
            with bcols[0]:
                if st.button("👁 View", key=f"view_{i}"):
                    st.session_state.last_query = q
                    st.session_state.selected_menu = None
                    st.rerun()
            with bcols[1]:
                docx = generate_docx(q, a)
                st.download_button("⬇️", data=docx, file_name=f"{slugify(q)}.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   key=f"download_{i}")
            with bcols[2]:
                if st.button("🗑 Delete", key=f"delete_{i}"):
                    st.session_state.chat_history.pop(idx)
                    st.session_state.selected_menu = None
                    st.rerun()

# --- Welcome Banner ---
st.markdown("""
    <div style='background-color: #f4f4f4; padding: 25px; border-radius: 10px; margin-bottom: 30px;'>
        <h1 style='color: #333; margin: 0;'>Welcome to CustomerBrief</h1>
        <p style='font-size: 16px; color: #444; margin-top: 10px;'>
            AI-powered insights for sharper sales conversations.
        </p>
    </div>
""", unsafe_allow_html=True)

# --- Query Input ---
user_query = st.text_input(
    label="", placeholder="Enter company name or market query...", label_visibility="collapsed"
)
search_clicked = st.button("🔍 Search")


# --- Handle Query with Detailed Business Overview ---
def process_with_openai(query):
    with st.spinner("🔍 Analyzing the business..."):
        try:
            # System prompt designed for up-to-date, detailed company analysis
            messages = [
                {"role": "system", "content": (
                    "You are CustomerBrief, an expert market analyst. "
                    "Provide a clear, structured, and insightful business analysis of any company using the most recent and relevant data available.\n\n"
                    "Your response must include:\n\n"
                    "1. 🏢 Company Overview\n"
                    "2. 💰 Financial Summary (revenue, profit, funding, etc.)\n"
                    "3. 🌍 Market & Industry Position\n"
                    "4. 📦 Import Activity\n"
                    "5. 🚢 Export Activity\n"
                    "6. 🌍 Global Presence & Office Locations\n"
                    "7. 🚛 Freight Forwarding History\n"
                    "8. 🔍 Competitive Landscape\n"
                    "9. 📈 Recent Developments or Strategic Moves\n"
                    "10. 🧠 Actionable Insights & Recommendations\n"
                    "11. 🔗 Source Links (insert relevant links within each section if available)\n\n"
                    "Be factual, neutral, and use bullet points or sub-sections for clarity. "
                    "Use the most current information and trends available at the time of analysis. "
                    "If some data is estimated or not publicly available, clearly mention that."
                )},
                {"role": "user", "content": query}
            ]

            response = client.chat.completions.create(
                model="gpt-4.1-2025-04-14",  # Use whichever model your system prefers
                messages=messages,
                temperature=0.4
            )

            result = response.choices[0].message.content.strip()

            st.success("✅ Analysis Complete")
            show_download_buttons(query, result)
            st.session_state.chat_history.append((query, result))

        except Exception as e:
            st.error(f"❌ OpenAI API Error: {e}")




# --- Run on Click ---
if search_clicked or (user_query and user_query != st.session_state.last_query):
    if user_query.strip():
        import re
        query = user_query.strip()

        # Heuristic: multiple separators or more than 2 capitalized "words"
        multi_company_separators = [" and ", ",", " & ", " vs ", "/", " versus "]
        has_multiple_separators = any(sep in query.lower() for sep in multi_company_separators)

        # Detect capitalized words (common in company names)
        company_like_keywords = re.findall(r"\b[A-Z][a-zA-Z&.\-']{2,}\b", query)
        likely_multiple_companies = len(company_like_keywords) > 2

        if has_multiple_separators or likely_multiple_companies:
            st.warning("⚠️ **Important Notice:** To ensure clarity and depth in analysis, our AI system is designed to evaluate one company at a time. "
                       "Please revise your query to reference a single organization for a precise and comprehensive report. 🏢")
        else:
            st.session_state.last_query = query
            process_with_openai(query)
    else:
        st.warning("Please enter a valid query.")




# --- Reload Last Output if Exists ---
if (
    st.session_state.last_query
    and not search_clicked
    and st.session_state.last_query not in [q for q, _ in st.session_state.chat_history[-1:]]
):
    # Only show saved last query result if NOT a fresh search submission
    for idx, (q, a) in enumerate(reversed(st.session_state.chat_history)):
        if q == st.session_state.last_query:
            st.markdown(f"### **User Query:** {q}")
            show_download_buttons(q, a, key_prefix=f"history_{idx}_{slugify(q)[:30]}")


