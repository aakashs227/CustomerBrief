import streamlit as st
import base64
import requests
from PIL import Image
from io import BytesIO
from docx import Document
import re

def clean_response(text):
    cleaned_text = re.sub(
        r"(🧠 Company Analysis\s*)(.*?)(?=\n\s*1\.\s*[🔍A-Z])",
        r"\1\n",
        text,
        flags=re.DOTALL
    )
    return cleaned_text.strip()

# --- Page Setup ---
st.set_page_config(
    page_title="CustomerBrief AI-powered insights for sharper sales conversations.",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Custom Styling ---
st.markdown("""
    <style>
        #MainMenu, header, footer {
            visibility: hidden;
            height: 0px;
        }

        section[data-testid="stSidebar"] {
            background-color: #002b5c;
        }

        .stButton button {
            background-color: white !important;
            color: black !important;
            border: 1px solid #999 !important;
        }

        .chat-history-item {
            color: yellow !important;
            font-size: 16px !important;
        }

        .top-right-logo {
            position: absolute;
            top: 10px;
            right: 30px;
            z-index: 1000;
        }

        .top-right-logo img {
            width: 150px;
            height: auto;
            display: block;
            border: none;
        }
    </style>
""", unsafe_allow_html=True)

# --- Load Logos ---
logo_path = "WORLDWIDE_Logo_7.png"  # Make sure the file name and case match exactly

try:
    if not os.path.exists(logo_path):
        raise FileNotFoundError(f"{logo_path} not found in current directory.")

    logo = Image.open(logo_path)
    buffered = BytesIO()
    logo.save(buffered, format="PNG")
    logo_base64 = base64.b64encode(buffered.getvalue()).decode()


# --- Initialize Session State ---
for key in ["chat_history", "last_query", "download_clicked", "share_clicked", "selected_menu"]:
    if key not in st.session_state:
        st.session_state[key] = [] if "history" in key else ""

# --- Sidebar ---
with st.sidebar:
    col_logo, col_button = st.columns([0.6, 0.4])
    with col_logo:
        st.markdown(
    f"""
    <div style="
        background-color: white;
        padding: 3px 0px 3px 0px;
        width: 100%;
        display: flex;
        justify-content: flex-start;
        align-items: center;
        box-sizing: border-box;
    ">
        <img src="data:image/png;base64,{logo_base64}" width="100" style="display: block;"/>
    </div>
    """,
    unsafe_allow_html=True
)



    st.markdown("---")
    st.markdown("""
        <h3 style="color: white; text-decoration: underline; margin-bottom: 10px;">
            Chat History
        </h3>
    """, unsafe_allow_html=True)

    for i, (q, a) in enumerate(reversed(st.session_state.chat_history)):
        index = len(st.session_state.chat_history) - 1 - i
        col1, col2 = st.columns([0.85, 0.15])
        with col1:
            st.markdown(
                f"<div class='chat-history-item'>• {q[:50]}</div>",
                unsafe_allow_html=True
            )
        with col2:
            if st.button("⋮", key=f"menu_{i}"):
                st.session_state.selected_menu = index

        if st.session_state.get("selected_menu") == index:
            button_cols = st.columns([1, 1, 1])
            with button_cols[0]:
                if st.button("👁 View", key=f"view_{i}", help="View"):
                    st.session_state.last_query = q
                    st.session_state.selected_menu = None
                    st.rerun()
            with button_cols[1]:
                doc_buffer = BytesIO()
                doc = Document()
                doc.add_heading("MIRA Company Analysis", level=1)
                doc.add_paragraph(f"Query: {q}")
                doc.add_paragraph("Response:")
                doc.add_paragraph(a)
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                st.download_button(
                    label="⬇️",
                    data=doc_buffer,
                    file_name=f"chat_history_{index+1}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_btn_{i}",
                    help="Download"
                )
            with button_cols[2]:
                if st.button("🗑 Delete", key=f"delete_{i}", help="Delete"):
                    st.session_state.chat_history.pop(index)
                    st.session_state.selected_menu = None   
                    st.rerun()

# --- Welcome Message ---
st.markdown("""
    <div style='background-color: #f4f4f4; padding: 25px; border-radius: 10px; margin-bottom: 30px;'>
        <h1 style='color: #333; margin: 0;'>Welcome to CustomerBrief</h1>
        <p style='font-size: 16px; color: #444; margin-top: 10px;'>
            AI-powered insights for sharper sales conversations.
        </p>
    </div>
""", unsafe_allow_html=True)

# --- Helper Functions ---
def slugify(text):
    return "".join(c if c.isalnum() else "_" for c in text.lower()).strip("_")

def generate_docx(query, response):
    doc = Document()
    doc.add_heading("MIRA Company Analysis", level=1)
    doc.add_paragraph(f"Query: {query}")
    doc.add_paragraph("Response:")
    doc.add_paragraph(response)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def contains_multiple_companies(query):
    delimiters = [",", "&", " and ", "/", " vs ", " versus ", ";"]
    count = sum([query.lower().count(d) for d in delimiters])
    return count >= 1

def show_download_buttons(query, response, key_prefix="main"):
    st.markdown("### 🧠 Company Analysis")
    if contains_multiple_companies(query):
        # Display the warning only once, not repeatedly
        if not hasattr(st.session_state, 'multiple_companies_warned'):
            st.warning("⚠️ Important Notice: To ensure clarity and depth in analysis, our AI system is designed to evaluate one company at a time. Please revise your query to reference a single organization for a precise and comprehensive report. 🏢")
            st.session_state.multiple_companies_warned = True  # Set flag to indicate warning has been shown
        st.write(response)
    else:
        file_name = f"{slugify(query)}.docx"
        docx_file_top = generate_docx(query, response)
        st.download_button(
            label="📥 Download Analysis",
            data=docx_file_top,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"{key_prefix}_download_top"
        )
        st.write(response)
        docx_file_bottom = generate_docx(query, response)
        st.download_button(
            label="📥 Download Analysis",
            data=docx_file_bottom,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"{key_prefix}_download_bottom"
        )

# --- Process Query ---
def process_query(query):
    with st.spinner("Processing..."):
        payload = {
            "model_name": "gpt-4.1-2025-04-14",
            "model_provider": "OpenAI",
            "messages": [query],
            "allow_search": True
        }
        try:
            # Send query to backend for processing
            response = requests.post("http://127.0.0.1:9999/chat", json=payload)
            data = response.json()
            if "error" in data:
                st.error(data["error"])  # Display error if response has an error
            else:
                st.success("Analysis complete.")  # Display success message

                # Show download and share buttons, along with the response
                show_download_buttons(query, data["response"], key_prefix="current")

                # Store the current query and response in session history
                st.session_state.chat_history.append((query, data["response"]))

                # Enable share button and download feature if the analysis is successful
                share_url = f"Check out the analysis for {query} at this link: [Company Analysis](http://127.0.0.1:9999/chat)"
                st.markdown(f"### Share this Analysis")
                st.markdown(f"[Share via Link]({share_url})")  # Shareable link to the analysis

        except Exception as e:
            # Handle exceptions, such as network errors
            st.error(f"❌ Error fetching response: {e}")

# --- Input Field ---
user_query = st.text_input(
    label="",
    placeholder="e.g., Enter Company Name or Query",
    label_visibility="collapsed",
    key="user_query"
)

# --- Search Button ---
search_clicked = st.button("🔍 Search", key="search_button")

# --- Trigger Query ---
if search_clicked or (user_query and user_query != st.session_state.last_query):
    if user_query:
        st.session_state.last_query = user_query
        process_query(user_query)
        st.session_state.input_temp = ""
        st.rerun()
    else:
        st.warning("Please enter a query before searching.")

# --- Show Saved Output if Needed ---
if st.session_state.last_query:
    for idx, (q, a) in enumerate(reversed(st.session_state.chat_history)):
        if q == st.session_state.last_query:
            st.markdown(f"### **User Query:** {q}")
            show_download_buttons(q, a, key_prefix=f"history_{idx}_{slugify(q)[:30]}")
