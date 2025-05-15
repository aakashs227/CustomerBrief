from io import BytesIO
from docx import Document
import uuid

def generate_docx_file(content: str) -> BytesIO:
    """
    Converts a string (company insights) into a .docx file and returns a BytesIO object.
    """
    doc = Document()
    doc.add_heading('Company Insights', 0)
    
    for line in content.split("\n"):
        doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_share_link(data: str) -> str:
    """
    Mocks generating a shareable link. Replace this with actual logic if needed.
    """
    unique_id = uuid.uuid4().hex
    # In real case, you'd store data in a database or pastebin-like service
    return f"https://mocksharelink.com/report/{unique_id}"
